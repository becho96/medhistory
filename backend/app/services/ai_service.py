import base64
import httpx
import json
from typing import Optional, Tuple
from io import BytesIO
from docx import Document as DocxDocument

from app.core.config import settings
from app.schemas.document import DocumentMetadata

# Модель читает PDF нативно (и текстовый слой, и сканы). Движок задан явно,
# чтобы OpenRouter не переключился на платный mistral-ocr.
PDF_PARSER_PLUGINS = [{"id": "file-parser", "pdf": {"engine": "native"}}]


class AIService:
    def __init__(self):
        self.api_key = settings.OPENROUTER_API_KEY
        self.model = settings.OPENROUTER_MODEL
        self.base_url = settings.OPENROUTER_BASE_URL
    
    async def analyze_document(self, file_bytes: bytes, file_type: str, filename: str) -> Tuple[DocumentMetadata, Optional[dict]]:
        """Analyze document and extract metadata using AI. Returns (metadata, usage)."""
        
        plugins = None

        try:
            # Handle different file types
            if file_type == 'pdf':
                # Send the PDF itself — the model reads both text layers and scans
                print("📄 Отправляем PDF в модель...")

                messages = self._build_pdf_messages(file_bytes, filename, include_full_text=True)
                plugins = PDF_PARSER_PLUGINS
                extracted_full_text = None
                full_text_source = "ai_vision_transcription"

            elif file_type == 'docx':
                print("📝 Извлекаем текст из DOCX...")
                text_content = self._extract_text_from_docx(file_bytes)

                if not text_content or len(text_content.strip()) < 20:
                    raise ValueError("DOCX содержит слишком мало текста для анализа.")

                print(f"  ✅ Извлечено {len(text_content)} символов текста")

                prompt = self._build_extraction_prompt(include_full_text=False)
                messages = [
                    {
                        "role": "user",
                        "content": f"{prompt}\n\nТекст медицинского документа:\n\n{text_content}"
                    }
                ]
                extracted_full_text = text_content
                full_text_source = "docx_text_extraction"

            elif file_type in ['jpg', 'jpeg', 'png']:
                # Use vision API for images
                print(f"🖼️ Обрабатываем изображение ({file_type})...")
                
                # Encode file to base64
                file_base64 = base64.b64encode(file_bytes).decode('utf-8')
                
                # Determine content type
                if file_type in ['jpg', 'jpeg']:
                    mime_type = 'image/jpeg'
                else:
                    mime_type = 'image/png'
                
                # Prepare vision message
                prompt = self._build_extraction_prompt(include_full_text=True)
                messages = [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{file_base64}"
                                }
                            }
                        ]
                    }
                ]
                extracted_full_text = None
                full_text_source = "ai_vision_transcription"
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            # Call OpenRouter API
            response_data = await self._call_openrouter(messages, plugins)

            # На очень длинных PDF транскрипция full_text может упереться в
            # max_tokens — ответ обрывается на середине JSON. Повторяем без неё:
            # классификация и summary важнее полного текста.
            if file_type == 'pdf' and self._is_truncated(response_data):
                print("✂️ Ответ обрезан по max_tokens — повтор без полной транскрипции")
                messages = self._build_pdf_messages(file_bytes, filename, include_full_text=False)
                response_data = await self._call_openrouter(messages, plugins)
                full_text_source = None

            # Parse response
            metadata = self._parse_ai_response(response_data)
            if extracted_full_text:
                metadata.full_text = extracted_full_text
                metadata.full_text_source = full_text_source
            elif metadata.full_text:
                metadata.full_text_source = metadata.full_text_source or full_text_source
            usage = self._extract_usage(response_data)
            return (metadata, usage)
            
        except Exception as e:
            print(f"❌ AI analysis failed: {str(e)}")
            # Return default metadata on error
            return (
                DocumentMetadata(
                    document_type="неизвестно",
                    confidence=0.0,
                    summary=f"Не удалось автоматически проанализировать документ: {str(e)}"
                ),
                None,
            )
    
    def _is_truncated(self, response_data: dict) -> bool:
        """True when the model ran out of output budget mid-answer."""
        return response_data["choices"][0].get("finish_reason") == "length"

    def _build_pdf_messages(self, file_bytes: bytes, filename: str, include_full_text: bool) -> list:
        """Message payload that hands the PDF itself to the model."""
        prompt = self._build_extraction_prompt(include_full_text=include_full_text)
        return [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    self._build_pdf_content_part(file_bytes, filename),
                ],
            }
        ]

    def _build_pdf_content_part(self, file_bytes: bytes, filename: str) -> dict:
        """Wrap raw PDF bytes into an OpenRouter file content part."""
        file_base64 = base64.b64encode(file_bytes).decode('utf-8')
        return {
            "type": "file",
            "file": {
                "filename": filename,
                "file_data": f"data:application/pdf;base64,{file_base64}",
            },
        }

    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract paragraphs and tables from a DOCX file as readable text."""
        try:
            docx_file = BytesIO(file_bytes)
            document = DocxDocument(docx_file)

            text_parts: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)

            for table_index, table in enumerate(document.tables, start=1):
                rows: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    text_parts.append(f"Таблица {table_index}:\n" + "\n".join(rows))

            return "\n\n".join(text_parts).strip()
        except Exception as e:
            print(f"❌ Ошибка извлечения текста из DOCX: {e}")
            raise ValueError(f"Не удалось извлечь текст из DOCX: {str(e)}")
    
    def _build_extraction_prompt(self, include_full_text: bool = False) -> str:
        content_fields = """
  "full_text": "подробная транскрипция всего видимого текста документа",
  "tables": [
    {
      "title": "название таблицы или null",
      "columns": ["колонка 1", "колонка 2"],
      "rows": [
        {"колонка 1": "значение", "колонка 2": "значение"}
      ]
    }
  ],""" if include_full_text else """
  "tables": [],"""

        content_rules = """
9. full_text - подробная транскрипция содержимого документа, а НЕ summary.
   Сохраняй исходные формулировки, порядок разделов, назначения, заключения,
   показатели и примечания. Не сокращай клинически значимые детали.
10. Если в документе есть таблицы, продублируй их данные в tables.
    Для результатов анализов сохраняй строки с названием показателя, значением,
    единицами измерения, референсами и отметками нормы/отклонения, если они есть.
11. full_text сохраняй на языке оригинального документа. Остальные классификационные
    поля возвращай на русском по правилам ниже.""" if include_full_text else """
9. tables верни пустым массивом []. Полный текст документа уже извлечен системным
   парсером и будет сохранен отдельно."""

        return f"""Проанализируй этот медицинский документ и извлеки следующую информацию в формате JSON.

# ОБЯЗАТЕЛЬНАЯ СТРУКТУРА JSON:

{{
  "document_type": "ОБЯЗАТЕЛЬНО - один из 5 типов",
  "document_subtype": "условно обязательно",
  "research_area": "только для инструментальных исследований",
  "specialties": ["массив специальностей"],
  "document_date": "YYYY-MM-DD",
  "patient_name": "Фамилия И.О.",
  "medical_facility": "название учреждения",
  "doctor_name": "ФИО врача, выдавшего документ, в формате Фамилия И.О., или null",
  "document_language": "ISO 639-1 код языка документа (например: ru, en, nl, de, fr)",
  "confidence": 0.95,
  "summary": "краткое содержание",
{content_fields}
  "orders": [
    {{
      "title": "УЗИ брюшной полости",
      "kind": "referral_research",
      "order_type": "instrumental",
      "target_document_type": "Инструментальное исследование",
      "target_document_subtype": "УЗИ",
      "target_research_area": "Брюшная полость",
      "target_specialty": null,
      "due_date": null,
      "due_after": {{"amount": 1, "unit": "month"}}
    }},
    {{
      "title": "Повторный приём кардиолога",
      "kind": "follow_up_appointment",
      "order_type": "consultation",
      "target_document_type": "Прием врача",
      "target_document_subtype": null,
      "target_research_area": null,
      "target_specialty": "Кардиология",
      "due_date": "2026-09-15",
      "due_after": null
    }}
  ]
}}

# ТИПЫ ДОКУМЕНТОВ (выбери ОДИН из 5):

1. "Прием врача" - амбулаторный прием, консультация, медицинское заключение врача
2. "Результаты анализа" - лабораторные анализы любого типа
3. "Инструментальное исследование" - УЗИ, МРТ, КТ, рентген, маммография
4. "Функциональная диагностика" - ЭКГ, ЭЭГ, холтер, спирометрия, ФГДС, колоноскопия
5. "Другое" - если не подходит ни под одну категорию

# ПРАВИЛА ЗАПОЛНЕНИЯ ПО ТИПАМ:

## "Прием врача":
- specialties: ["Кардиология", "Терапия"] - ОБЯЗАТЕЛЬНО массив (минимум 1)
- document_subtype: null
- research_area: null

Справочник специальностей: Терапия, Кардиология, Неврология, Гастроэнтерология, Эндокринология, Урология, Оториноларингология (ЛОР), Офтальмология, Травматология и ортопедия, Дерматология, Инфекционные болезни, Ревматология, Пульмонология, Нефрология, Онкология, Хирургия, Акушерство и гинекология, Педиатрия, Психиатрия, Другая специальность

## "Результаты анализа":
- document_subtype: "Общий анализ крови" - ОБЯЗАТЕЛЬНО
- specialties: null
- research_area: null

Подтипы: Общий анализ крови, Биохимический анализ крови, Общий анализ мочи, Анализ кала, Бактериологический анализ, Серологический анализ, Гистологическое исследование, Цитологическое исследование, Иммунологический анализ, Гормональный анализ, Генетический анализ, Другой анализ

## "Инструментальное исследование":
- document_subtype: "УЗИ" - ОБЯЗАТЕЛЬНО
- research_area: "Брюшная полость" - рекомендуется
- specialties: null

Подтипы: УЗИ, МРТ, КТ, Рентген, Маммография, Флюорография, Ангиография, Другое исследование

Области для УЗИ: Брюшная полость, Щитовидная железа, Молочные железы, Органы малого таза, Сердце (ЭхоКГ), Сосуды, Суставы, Мягкие ткани, Другая область

Области для МРТ/КТ: Головной мозг, Позвоночник, Брюшная полость, Грудная клетка, Суставы, Мягкие ткани, Сосуды, Другая область

Области для Рентген: Грудная клетка, Позвоночник, Суставы, Кости конечностей, Черепа, Другая область

## "Функциональная диагностика":
- document_subtype: "ЭКГ (электрокардиография)" - ОБЯЗАТЕЛЬНО
- specialties: null
- research_area: null

Подтипы: ЭКГ (электрокардиография), ЭЭГ (электроэнцефалография), Холтер-мониторирование, Суточный мониторинг АД, Спирометрия, Велоэргометрия, ФГДС (фиброгастродуоденоскопия), Колоноскопия, Бронхоскопия, Другая функциональная диагностика

## "Другое":
- Все дополнительные параметры опциональны
- confidence должна быть низкой (<0.5)

# ОБЩИЕ ПРАВИЛА:

1. document_type - ВСЕГДА ОБЯЗАТЕЛЬНО, используй точное название из списка выше
2. Дата документа - это дата процедуры/приема, НЕ дата печати
3. Формат ФИО: "Иванов И.И." (с точками). doctor_name - ФИО врача, который вёл приём/
   выдал заключение или направление (лечащий/консультирующий врач), если оно указано;
   иначе null. Не путай с пациентом (patient_name).
4. Специальности (specialties) - СУЩЕСТВИТЕЛЬНЫЕ: "Кардиология", а НЕ "Кардиолог"
5. Если информация не найдена - используй null
6. confidence (0.0-1.0) - твоя уверенность в классификации
7. document_language - ISO 639-1 код языка оригинала ("ru", "en", "nl", "de" и т.д.)
8. orders - список назначений/направлений врача, если они явно указаны в документе.
   Заполняй его прежде всего для "Прием врача". Если назначений нет, верни [].
   Каждому назначению задай "kind" (один из):
   - "referral_research" - направление на анализ/инструментальное/функциональное исследование;
     order_type: "lab", "instrumental" или "functional";
     target_document_type: "Результаты анализа", "Инструментальное исследование" или "Функциональная диагностика";
     target_document_subtype и target_research_area заполняй, когда это можно понять из текста.
   - "referral_specialist" - направление к врачу ДРУГОЙ специальности (консультация);
     order_type: "consultation"; target_document_type: "Прием врача";
     target_specialty - название специальности из справочника выше (существительное, напр. "Неврология").
   - "follow_up_appointment" - повторный приём у того же/лечащего врача;
     order_type: "consultation"; target_document_type: "Прием врача";
     target_specialty - специальность приёма, если известна.
   Не включай лекарства и общие рекомендации по режиму, питанию, образу жизни.
   Срок (когда явно указан): "due_date" - абсолютная дата "YYYY-MM-DD", ТОЛЬКО если она прямо
   названа в документе. Для относительных формулировок ("через 3 месяца", "через 2 недели")
   заполняй "due_after" объектом {{"amount": <число>, "unit": "day"|"week"|"month"|"year"}}, а
   "due_date" оставляй null - точную дату вычислит система от даты приёма. Если срок не указан -
   оба поля null.
{content_rules}

# ИНОСТРАННЫЕ ДОКУМЕНТЫ:
Документ может быть на любом языке (русский, английский, голландский, немецкий и др.).
Классифицируй документ по содержанию, а не по языку. Значения полей document_type,
document_subtype, research_area, specialties, summary, patient_name и medical_facility
всегда возвращай на русском языке (переведи при необходимости) — используй точные
строки из списков выше для document_type, document_subtype и specialties.

# ВАЖНО:
- Отвечай ТОЛЬКО валидным JSON
- Без дополнительного текста
- Проверь, что структура соответствует типу документа"""
    
    async def _call_openrouter(self, messages: list, plugins: Optional[list] = None) -> dict:
        """Make API call to OpenRouter"""

        payload = {
            "model": self.model,
            "messages": messages,
            "max_tokens": settings.OPENROUTER_MAX_TOKENS,
            "temperature": 0.1  # Low temperature for consistent extraction
        }
        if plugins:
            payload["plugins"] = plugins

        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost:8000",
            "X-Title": "MedHistory"
        }
        
        # PII-safe logging: model + payload shape only, no content
        print(f"🔍 OpenRouter Request: model={self.model}, messages={len(messages)}")

        # Транскрипция многостраничного скана легко переваливает за минуту
        async with httpx.AsyncClient(timeout=180.0) as client:
            try:
                response = await client.post(
                    self.base_url,
                    headers=headers,
                    json=payload
                )

                print(f"📥 OpenRouter Response: status={response.status_code}")
                response.raise_for_status()
                return response.json()
            except httpx.HTTPStatusError as e:
                print(f"❌ OpenRouter HTTP Error: status={e.response.status_code}")
                raise
            except Exception as e:
                print(f"❌ OpenRouter Request Error: {type(e).__name__}")
                raise
    
    def _extract_usage(self, response_data: dict) -> Optional[dict]:
        """Extract token usage from OpenRouter response."""
        usage = response_data.get("usage")
        if not usage:
            return None
        return {
            "prompt_tokens": usage.get("prompt_tokens"),
            "completion_tokens": usage.get("completion_tokens"),
            "total_tokens": usage.get("total_tokens"),
        }
    
    def _parse_ai_response(self, response_data: dict) -> DocumentMetadata:
        """Parse AI response and create DocumentMetadata"""
        
        try:
            # Extract content from response
            content = response_data["choices"][0]["message"]["content"]
            
            # Parse JSON from content
            # Sometimes the model wraps JSON in markdown code blocks
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0].strip()
            elif "```" in content:
                content = content.split("```")[1].split("```")[0].strip()
            
            data = json.loads(content)
            tables = data.get("tables") if isinstance(data.get("tables"), list) else []
            
            # Create DocumentMetadata object with new structure
            metadata = DocumentMetadata(
                # PostgreSQL fields
                document_type=data.get("document_type", "Другое"),
                document_date=data.get("document_date"),
                patient_name=data.get("patient_name"),
                medical_facility=data.get("medical_facility"),
                doctor_name=data.get("doctor_name"),
                
                # MongoDB classification fields
                document_subtype=data.get("document_subtype"),
                research_area=data.get("research_area"),
                specialties=data.get("specialties"),
                document_language=data.get("document_language", "ru"),
                confidence=data.get("confidence", 0.5),
                
                # MongoDB extracted_data fields
                summary=data.get("summary"),
                full_text=data.get("full_text") if isinstance(data.get("full_text"), str) else None,
                tables=[t for t in tables if isinstance(t, dict)],
                orders=data.get("orders") if isinstance(data.get("orders"), list) else []
            )
            
            return metadata
            
        except Exception as e:
            print(f"❌ Error parsing AI response: {type(e).__name__}")
            raise

    async def generate_report_content(self, documents: list, filters: dict) -> str:
        """Generate report content using AI"""
        
        # Prepare context from documents
        context = self._prepare_report_context(documents)
        
        prompt = f"""Создай подробный медицинский отчёт для пациента на основе следующих документов:

{context}

Отчёт должен включать:
1. Общую информацию о пациенте
2. Хронологическую временную линию событий
3. Ключевые результаты анализов и исследований
4. Назначенное лечение и его динамику
5. Резюме текущего состояния здоровья

Формат отчёта должен быть структурированным с четкими заголовками.
Пиши на русском языке.
Будь профессиональным и детальным."""
        
        messages = [
            {
                "role": "user",
                "content": prompt
            }
        ]
        
        response_data = await self._call_openrouter(messages)
        content = response_data["choices"][0]["message"]["content"]
        
        return content
    
    def _prepare_report_context(self, documents: list) -> str:
        """Prepare context from documents for report generation"""
        
        context_parts = []
        
        for doc in documents:
            doc_info = f"""
Дата: {doc.document_date or 'не указана'}
Тип: {doc.document_type or 'не указан'}
Пациент: {doc.patient_name or 'не указан'}
Учреждение: {doc.medical_facility or 'не указано'}
Файл: {doc.original_filename}
"""
            context_parts.append(doc_info)
        
        return "\n---\n".join(context_parts)

    async def extract_lab_results(self, file_bytes: bytes, file_type: str, filename: str) -> dict:
        """Extract laboratory results using a specialized prompt.

        Returns a dict with key "lab_results": list of standardized entries.
        """
        # Build labs prompt
        prompt = self._build_labs_extraction_prompt()

        # Prepare messages similarly to analyze_document
        plugins = None
        if file_type == 'pdf':
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        self._build_pdf_content_part(file_bytes, filename),
                    ],
                }
            ]
            plugins = PDF_PARSER_PLUGINS
        elif file_type == 'docx':
            text_content = self._extract_text_from_docx(file_bytes)
            messages = [
                {
                    "role": "user",
                    "content": f"{prompt}\n\nТекст медицинского документа:\n\n{text_content}"
                }
            ]
        elif file_type in ['jpg', 'jpeg', 'png']:
            file_base64 = base64.b64encode(file_bytes).decode('utf-8')
            mime_type = 'image/jpeg' if file_type in ['jpg', 'jpeg'] else 'image/png'
            messages = [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{file_base64}"}},
                    ],
                }
            ]
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        response_data = await self._call_openrouter(messages, plugins)
        return self._parse_labs_response(response_data)

    def _build_labs_extraction_prompt(self) -> str:
        return """Определи, есть ли в документе результаты лабораторных анализов. Если да, извлеки их в JSON по стандартизованной схеме.

Документ может быть на любом языке (русский, английский, голландский, немецкий и др.).
Не возвращай пустой результат только потому, что документ не на русском — извлеки анализы
и переведи названия и единицы измерения на русский язык.

Стандартизованный формат:
{
  "lab_results": [
    {
      "test_name": "название анализа на русском (например: гемоглобин, глюкоза)",
      "value": "числовое значение как строка",
      "unit": "единицы измерения на русском (например: г/л, ммоль/л, %)",
      "reference_range": "референсный диапазон как строка (например: 120-160)",
      "flag": "L|N|H|A"  // L ниже нормы, N норма, H выше нормы, A абнормально/критично
    }
  ]
}

Правила:
- Если анализов в документе действительно нет, верни {"lab_results": []}
- test_name всегда на русском — если документ на другом языке, переведи название
  (Hemoglobine/Hemoglobin → Гемоглобин, Glucose → Глюкоза, Cholesterol → Холестерин и т.д.)
- value оставляй строкой, не округляй
- unit переводи в общепринятый русский вид, сохраняя смысл единицы из документа:
  * g/L → г/л, g/dL → г/дл, mmol/L → ммоль/л, mg/dL → мг/дл, µmol/L → мкмоль/л
  * IU/L → МЕ/л, 10^9/L → ×10⁹/л, 10^12/L → ×10¹²/л
  * Проценты (%) и pH оставляй как есть
  * Если единиц нет в документе, используй null
- ВАЖНО про единицы:
  * Разные единицы измерения одного анализа (например, процентные и абсолютные) - это РАЗНЫЕ биомаркеры
  * Не смешивай единицы: если в документе "Лимфоциты 42%" и "Лимфоциты 2.5 10*9/л", это две отдельные записи
- Если референса нет, используй null
- Определи flag исходя из референсного диапазона, если он указан
- Отвечай ТОЛЬКО валидным JSON

Примеры правильного извлечения:
- "Лимфоциты 42%" → {"test_name": "Лимфоциты", "value": "42", "unit": "%", ...}
- "Лимфоциты 2.5 10*9/л" → {"test_name": "Лимфоциты", "value": "2.5", "unit": "10*9/л", ...}
- "Гемоглобин 145 г/л" → {"test_name": "Гемоглобин", "value": "145", "unit": "г/л", ...}
- "Hemoglobine 8.7 mmol/L" → {"test_name": "Гемоглобин", "value": "8.7", "unit": "ммоль/л", ...}
- "Glucose 5.4 mmol/L (3.9-5.6)" → {"test_name": "Глюкоза", "value": "5.4", "unit": "ммоль/л", "reference_range": "3.9-5.6", "flag": "N"}
- "Cholesterol 220 mg/dL" → {"test_name": "Холестерин", "value": "220", "unit": "мг/дл", ...}
"""

    def _parse_labs_response(self, response_data: dict) -> dict:
        # A truncated response (output hit max_tokens) yields invalid JSON. We
        # must NOT swallow that as an empty result — an empty lab list would
        # look like a successfully processed document with no analytes. Surface
        # it so the caller can log/retry. A genuine "no labs" case still comes
        # back as valid JSON ({"lab_results": []}) and is handled normally.
        choice = response_data["choices"][0]
        if choice.get("finish_reason") == "length":
            raise ValueError(
                "Labs extraction response truncated (finish_reason=length); "
                "raise OPENROUTER_MAX_TOKENS"
            )
        try:
            content = choice["message"]["content"]
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0].strip()
            elif "```" in content:
                content = content.split("```")[1].split("```")[0].strip()
            data = json.loads(content)

            # Basic normalization
            lab_results = data.get("lab_results") or []
            normalized = []
            for item in lab_results:
                normalized.append({
                    "test_name": item.get("test_name"),
                    "value": item.get("value"),
                    "unit": item.get("unit"),
                    "reference_range": item.get("reference_range"),
                    "flag": item.get("flag"),
                })
            return {
                "lab_results": normalized,
                "usage": self._extract_usage(response_data),
            }
        except json.JSONDecodeError as e:
            # Malformed/partial JSON — treat as a failure, not an empty result.
            print(f"❌ Error parsing labs response: {type(e).__name__}")
            raise ValueError("Labs extraction returned invalid JSON") from e

# Create global instance
ai_service = AIService()
